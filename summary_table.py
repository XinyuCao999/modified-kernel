"""
Summary tables for the reviewer response.

For every case / acquisition function / method, this reads the 50-repetition BO
results and reports the distance-to-optimum of the LAST sampling point:
mean and standard deviation across the 50 repetitions.

    distance = |reference_max - best_so_far|
    best_so_far = running maximum of the sampled objective values
    (identical definition to plot_bo_comparison in BO_base.py)

Two tables are produced (printed to the console AND written to a Word file):

    Table 1 - main summary: all three cases, four methods,
              columns = UCB (beta=2) and EI, cells = mean +- std.
    Table 2 - Case 3 UCB beta sensitivity: four methods,
              columns = beta in {2, 2.5, 3, 3.5, 4}, cells = mean +- std.

Run:
    python3 summary_table.py
"""
import os
import sys
import types
import numpy as np

# ---------------------------------------------------------------------------
# The .npy files are pickled dicts whose "machanistic_fun" value is a function
# reference in the case-specific "*_process" module (importing it pulls in
# torch). We only need the numeric "output_data_all" array, so we stub any such
# module on demand: np.load then succeeds without executing the real process
# file and without importing torch. The stored numbers are untouched.
# ---------------------------------------------------------------------------
class _StubModule(types.ModuleType):
    def __getattr__(self, name):
        return lambda *a, **k: None

def _load_dict(path):
    while True:
        try:
            return np.load(path, allow_pickle=True).item()
        except ModuleNotFoundError as e:
            sys.modules[e.name] = _StubModule(e.name)

# ---------------------------------------------------------------------------
# Configuration.
# ---------------------------------------------------------------------------
DATE = "0726"
BETA_MAIN = 2                       # UCB beta used in the main summary table
CASE3_BETAS = [2, 2.5, 3, 3.5, 4]   # UCB betas for the Case 3 sensitivity table

CASES = [
    ("Case 1 - Van de Vusse",      "Case1-vdv/BO",               0.7888299971987257),
    ("Case 2 - Williams-Otto",     "Case2-william otto/BO",      191.02),
    ("Case 3 - Schotten-Baumann",  "Case3-schotten baumann/BO",  0.89351549424855),
]
CASE3 = CASES[2]

# method label -> filename stem
METHODS = [
    ("zero-prior",     "RBF_zeroprior"),
    ("mech-prior",     "RBF_nonzeroprior"),
    ("multi-fidelity", "multi_fidelity"),
    ("modified",       "reconstructed"),
]


def _beta_tag(beta):
    """Match the filenames: integer betas have no decimals (beta2, beta3)."""
    return str(int(beta)) if float(beta).is_integer() else str(beta)


def final_error_stats(folder, stem, acq, ref, beta=BETA_MAIN):
    """Return (n_rep, mean, std) of the last-point distance-to-optimum, or None."""
    if acq == "UCB":
        fname = f"{stem}_{DATE}_UCB_beta{_beta_tag(beta)}.npy"
    else:
        fname = f"{stem}_{DATE}_EI.npy"
    path = os.path.join(folder, fname)
    if not os.path.exists(path):
        return None
    d = _load_dict(path)
    y = np.array(d["output_data_all"], dtype=float)     # [n_rep, n_points]
    best_so_far = np.maximum.accumulate(y, axis=1)
    final_error = np.abs(ref - best_so_far[:, -1])      # last point per rep
    return y.shape[0], final_error.mean(), final_error.std()


def cell(stats):
    if stats is None:
        return "[missing]"
    _, m, s = stats
    return f"{m:.3e} +/- {s:.3e}"


# ---------------------------------------------------------------------------
# Regret reduction of "modified" relative to the strongest (best) baseline.
# For each configuration the strongest baseline is the one with the smallest
# mean final regret among {zero-prior, mech-prior, multi-fidelity}; the
# reported reduction is (baseline - modified) / baseline. "up to" = the maximum
# reduction across the evaluated configurations of a case.
# ---------------------------------------------------------------------------
def regret_reduction(case, betas):
    """Yield (config_label, best_baseline, baseline_mean, modified_mean, reduction%)."""
    _, folder, ref = case
    baselines = [(l, s) for l, s in METHODS if s != "reconstructed"]
    configs = [("EI", None)] + [("UCB", b) for b in betas]
    for acq, beta in configs:
        b = beta if beta is not None else BETA_MAIN
        base_means = {l: final_error_stats(folder, s, acq, ref, beta=b)[1] for l, s in baselines}
        best_label = min(base_means, key=base_means.get)
        base_mean = base_means[best_label]
        mod_mean = final_error_stats(folder, "reconstructed", acq, ref, beta=b)[1]
        red = (base_mean - mod_mean) / base_mean * 100.0
        label = "EI" if acq == "EI" else f"UCB beta={_beta_tag(b)}"
        yield label, best_label, base_mean, mod_mean, red


def print_regret_reduction(case, betas):
    case_name = case[0]
    rows = list(regret_reduction(case, betas))
    header = ["Config", "strongest baseline", "baseline regret", "modified", "reduction %"]
    widths = [16, 20, 15, 12, 12]
    print("\n" + "=" * (sum(widths) + len(widths) * 2))
    print(f"  Regret reduction of 'modified' vs the strongest baseline - {case_name}")
    print("=" * (sum(widths) + len(widths) * 2))
    print("  ".join(h.ljust(widths[i]) for i, h in enumerate(header)))
    best = max(rows, key=lambda r: r[4])
    for label, bl, bm, mm, red in rows:
        print("  ".join(str(v).ljust(widths[i]) for i, v in enumerate(
            [label, bl, f"{bm:.3e}", f"{mm:.3e}", f"{red:.1f}%"])))
    print("-" * (sum(widths) + len(widths) * 2))
    print(f"  => up to {best[4]:.1f}% reduction  (config {best[0]}, "
          f"vs {best[1]}: {best[2]:.3e} -> {best[3]:.3e})")


# ---------------------------------------------------------------------------
# Build table contents as lists of rows (list of strings).
# ---------------------------------------------------------------------------
def build_table1():
    header = ["Case", "Method", "n", f"UCB (beta={BETA_MAIN}) mean +/- std", "EI mean +/- std"]
    rows = []
    for case_name, folder, ref in CASES:
        for i, (label, stem) in enumerate(METHODS):
            ucb = final_error_stats(folder, stem, "UCB", ref)
            ei = final_error_stats(folder, stem, "EI", ref)
            n = (ucb or ei or (0,))[0]
            rows.append([case_name if i == 0 else "", label, str(n), cell(ucb), cell(ei)])
    return header, rows


def build_table2():
    case_name, folder, ref = CASE3
    header = ["Method", "n"] + [f"UCB beta={_beta_tag(b)}" for b in CASE3_BETAS]
    rows = []
    for label, stem in METHODS:
        stats = [final_error_stats(folder, stem, "UCB", ref, beta=b) for b in CASE3_BETAS]
        n = next((s[0] for s in stats if s is not None), 0)
        rows.append([label, str(n)] + [cell(s) for s in stats])
    return header, rows


# ---------------------------------------------------------------------------
# Console printing.
# ---------------------------------------------------------------------------
def print_table(title, header, rows, subtitle=None):
    widths = [max(len(header[c]), *(len(r[c]) for r in rows)) for c in range(len(header))]
    def line(cells):
        return "  ".join(str(c).ljust(widths[i]) for i, c in enumerate(cells))
    bar = "=" * len(line(header))
    print("\n" + bar)
    print("  " + title)
    if subtitle:
        print("  " + subtitle)
    print(bar)
    print(line(header))
    print("-" * len(bar))
    for r in rows:
        print(line(r))
    print(bar)


# ---------------------------------------------------------------------------
# Word export.
# ---------------------------------------------------------------------------
def write_word(path, t1, t2):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Summary of BO performance across 50 repetitions", level=1)
    doc.add_paragraph(
        "Distance to optimum at the final sampling point, reported as "
        "mean +/- standard deviation over 50 independent repetitions. "
        "Distance = |reference optimum - best-so-far objective value|. "
        "Lower is better; the best method in each column is expected to have "
        "both the smallest mean and a small standard deviation."
    )

    def add_table(title, header, rows):
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=len(header))
        table.style = "Light Grid Accent 1"
        for j, h in enumerate(header):
            run = table.rows[0].cells[j].paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(9)
        for r in rows:
            cells = table.add_row().cells
            for j, val in enumerate(r):
                p = cells[j].paragraphs[0]
                run = p.add_run(val.replace("+/-", "±"))
                run.font.size = Pt(9)

    h1, r1 = t1
    add_table("Table 1. Final-point distance to optimum (all cases, UCB and EI).", h1, r1)

    doc.add_paragraph()
    h2, r2 = t2
    add_table("Table 2. Case 3 (Schotten-Baumann) UCB beta sensitivity.", h2, r2)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Reference optima: ").bold = True
    p.add_run("; ".join(f"{name} = {ref}" for name, _, ref in CASES) + ".")

    doc.save(path)
    return path


def main():
    t1 = build_table1()
    t2 = build_table2()

    print_table(
        "Table 1. Final-point distance to optimum (50 repetitions)",
        *t1, subtitle="distance = |reference_max - best_so_far|",
    )
    print_table(
        "Table 2. Case 3 (Schotten-Baumann) UCB beta sensitivity (50 repetitions)",
        *t2, subtitle="distance = |reference_max - best_so_far|",
    )

    # Regret reduction of "modified" vs the strongest baseline for Case 3.
    print_regret_reduction(CASE3, CASE3_BETAS)

    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "summary_tables.docx")
    write_word(out, t1, t2)
    print(f"\nWord file written to: {out}")


if __name__ == "__main__":
    main()
